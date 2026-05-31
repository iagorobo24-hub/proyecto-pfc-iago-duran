"""
Script DEFINITIVO para aplicar TODAS las correcciones a la MEMORIA_PFC_V1.docx
Según el PROMPT MAESTRO — Mejora de la Memoria del TFC
"""

import sys
sys.path.insert(0, '/tmp/docx-env/lib/python3.11/site-packages')

from docx import Document

INPUT_PATH = "/home/abu/github_repos/proyecto-pfc-iago-duran/proyecto-fin-ciclo/desarrollo-entrega-final/MEMORIA_PFC_V1.docx"
OUTPUT_PATH = "/home/abu/github_repos/proyecto-pfc-iago-duran/proyecto-fin-ciclo/desarrollo-entrega-final/MEMORIA_PFC_V3.docx"

doc = Document(INPUT_PATH)
cambios = 0

def aplicar_cambio(para, antiguo, nuevo, descripcion=""):
    global cambios
    if antiguo in para.text:
        para.text = para.text.replace(antiguo, nuevo)
        cambios += 1
        if descripcion:
            print(f"  [{descripcion}]")
        return True
    return False

print("\nTAREA 1: Cifras de productos")
print("-"*40)
for para in doc.paragraphs:
    aplicar_cambio(para, "400K+ productos", "~3.000 productos (migrados desde ~75.000 de Firestore)", "T1")
    aplicar_cambio(para, "400.000+", "~3.000 (de ~75.000 en Firestore)", "T1")

print("\nTAREA 2: Firebase → Supabase")
print("-"*40)
for para in doc.paragraphs:
    aplicar_cambio(para, "LoginPage → Firebase Google OAuth", "LoginPage → Supabase Google OAuth", "T2")
    aplicar_cambio(para, "onAuthStateChanged", "onAuthStateChange (Supabase)", "T2")
    aplicar_cambio(para, "Auth: Firebase Auth\n", "Auth: Supabase Auth\n", "T2")
    aplicar_cambio(para, "Base de datos: Firestore (→ Supabase en)", "Base de datos: Supabase (PostgreSQL)", "T2")
    aplicar_cambio(para, "Actual: Firestore con sincronización", "Anterior: Firestore (fase inicial)", "T2")
    aplicar_cambio(para, "Este documento describe el modelo actual en Firestore.", 
                   "Este documento describe el modelo anterior (Firestore fue usado en la fase inicial).", "T2")
    aplicar_cambio(para, "Modelo de datos en Supabase (futuro)", "Modelo actual — Supabase (PostgreSQL)", "T2")
    aplicar_cambio(para, "Estado: Migración en desarrollo", "", "T2")
    aplicar_cambio(para, "Firebase Security Rules", "Supabase RLS (Row Level Security)", "T2")
    aplicar_cambio(para, "Firebase Auth → Spark", "Supabase Auth → Free", "T2")
    aplicar_cambio(para, "Firestore → Spark", "Supabase → Free", "T2")
    aplicar_cambio(para, "Firebase Auth (Google Sign-In)", "Supabase Auth (Google Sign-In)", "T2")

print("\nTAREA 3: Frase rota")
print("-"*40)
for para in doc.paragraphs:
    if "Este proyecto 15+ herramientas de IA" in para.text:
        para.text = "En este proyecto se utilizaron más de 15 herramientas de IA distintas, desde Claude Web en marzo de 2026 hasta Hermes Agent en mayo de 2026. La metodología fue evolucionando conforme aparecían nuevas opciones:"
        cambios += 1
        print("  [Frase completada]")

print("\nTAREA 4: Eliminar Capítulo 02")
print("-"*40)
for para in doc.paragraphs:
    if para.text.strip() in ["*Capítulo 02 — Estado del Arte*", "Capítulo 02 — Estado del Arte"]:
        para.text = ""
        cambios += 1
        print("  [Eliminado vestigio]")

print("\nTAREA 5: Estructura del documento")
print("-"*40)
estructura_correcta = """Estructura del documento

Este TFC se documenta en 10 capítulos más una sección de anexos:

1. Introducción — Presentación del proyecto, objetivos y contexto
2. Análisis de la empresa — Sonepar: estructura, productos y contexto de las prácticas
3. Análisis de requisitos — Problemas identificados, requisitos funcionales y no funcionales
4. Diseño técnico — Arquitectura, stack tecnológico, UI/UX y modelo de datos
5. Contexto de la IA — Estado del arte de la IA generativa en desarrollo web
6. Herramientas de IA utilizadas — Catálogo de 15 herramientas con descripción y uso en el proyecto
7. Proceso de desarrollo — Metodología de trabajo con IA y cronología de las fases
8. Manuales de uso — Guías de los módulos principales
9. Resultados — Evaluación cualitativa y cuantitativa
10. Conclusiones y líneas futuras — Valoración final e impacto educativo
Anexos — Fichas técnicas completas de las herramientas de IA y tablas de referencia"""

encontrado = False
for i, para in enumerate(doc.paragraphs):
    if "Estructura del documento" in para.text and not encontrado:
        para.text = estructura_correcta
        cambios += 1
        encontrado = True
        # Vaciar párrafos siguientes hasta el próximo capítulo
        j = i + 1
        while j < len(doc.paragraphs):
            if "2. ANÁLISIS" in doc.paragraphs[j].text or "2.Análisis" in doc.paragraphs[j].text:
                break
            doc.paragraphs[j].text = ""
            j += 1
        print("  [Estructura corregida]")

print("\nTAREA 7: Contexto 9.2")
print("-"*40)
contexto_92 = """Las métricas de uso de esta sección son de entorno de desarrollo: un único usuario (yo), sesiones de prueba, y KPIs con datos simulados para verificar que el sistema funciona correctamente. No hay usuarios reales porque el proyecto no se lanzó públicamente durante el ciclo. Lo que sí son reales son las métricas de rendimiento (Lighthouse), el tamaño del bundle, los datos del catálogo y los resultados de los tests E2E. El valor del proyecto como TFC no es la escala de uso, sino demostrar que la arquitectura funciona y que el proceso está documentado."""

for para in doc.paragraphs:
    if "9.2-Resultados Cuantitativos" in para.text or "9.2 Resultados Cuantitativos" in para.text:
        para.text = para.text + "\n\n" + contexto_92
        cambios += 1
        print("  [Contexto añadido]")
        break

print("\nTAREA 8: Autocrítica capítulo 10")
print("-"*40)
autocritica = """Los tres objetivos no cumplidos —tests unitarios, CI/CD completo y validación con usuarios reales— tienen en común el mismo motivo: prioricé avanzar en funcionalidad sobre consolidar lo que ya funcionaba. Es una decisión que entiendo pero que no volvería a tomar. Sin tests, cada refactor era arriesgado. Sin validación externa, no sé si lo que construí resuelve el problema de verdad o solo lo resuelve desde mi perspectiva. Si repitiera el proyecto, implementaría una suite básica de tests E2E desde la fase 2 y buscaría al menos una sesión de feedback con técnicos de la delegación antes de cerrar el desarrollo."""

for i, para in enumerate(doc.paragraphs):
    if "objetivos no cumplidos" in para.text.lower() and "secundarios" in para.text.lower():
        para.text = "Objetivos no cumplidos\n\n" + autocritica
        cambios += 1
        print("  [Autocrítica añadida]")
        # Vaciar siguientes párrafos relacionados
        for j in range(i+1, min(i+3, len(doc.paragraphs))):
            if "secundarios" in doc.paragraphs[j].text.lower():
                doc.paragraphs[j].text = ""
        break

print("\nTAREA 9: Claude Web")
print("-"*40)
claude_completo = """Ahora uso Claude Web principalmente para razonar sobre arquitectura y planificar features nuevas. Le puedo mandar archivos específicos o documentación que explica cómo funciona "X", y él me genera una lista de tareas y componentes que hay que crear para cumplir el objetivo. Luego ejecuto esas tareas con Windsurf u OpenCode.

Por ejemplo, si quisiera incluir una herramienta más en la suite, primero le paso a Claude los manuales existentes, la estructura de carpetas del proyecto y una descripción de la nueva funcionalidad. Él analiza los patrones que ya uso, identifica qué componentes puedo reutilizar y qué hay que crear desde cero, y me devuelve una lista ordenada de tareas con los archivos que tocará modificar. Luego cojo esa lista y la voy ejecutando paso a paso con el IDE."""

for para in doc.paragraphs:
    if "le puedo mandar archivos específicos" in para.text and "Por ejemplo, en caso de que quisiera" in para.text:
        para.text = claude_completo
        cambios += 1
        print("  [Ficha Claude completada]")
        break

print("\nTAREA 10: Anglicismos")
print("-"*40)
for para in doc.paragraphs:
    aplicar_cambio(para, "Alternativas considered", "Alternativas consideradas", "T10")
    aplicar_cambio(para, "depois", "después", "T10")

print("\n" + "="*80)
print(f"TOTALES: {cambios} cambios aplicados")
print(f"Guardando: {OUTPUT_PATH}")
print("="*80)

doc.save(OUTPUT_PATH)
print("\n✅ Documento V3 generado exitosamente")