"""
Script para aplicar correcciones a la MEMORIA_PFC_V1.docx
Según el PROMPT MAESTRO — Mejora de la Memoria del TFC
"""

import sys
sys.path.insert(0, '/tmp/docx-env/lib/python3.11/site-packages')

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Ruta del documento
INPUT_PATH = "/home/abu/github_repos/proyecto-pfc-iago-duran/proyecto-fin-ciclo/desarrollo-entrega-final/MEMORIA_PFC_V1.docx"
OUTPUT_PATH = "/home/abu/github_repos/proyecto-pfc-iago-duran/proyecto-fin-ciclo/desarrollo-entrega-final/MEMORIA_PFC_V2.docx"

# Cargar documento
doc = Document(INPUT_PATH)

# Contador de cambios
cambios_realizados = 0

def reemplazar_texto_en_parrafo(para, texto_antiguo, texto_nuevo):
    """Reemplaza texto en un párrafo manteniendo el formato"""
    if texto_antiguo in para.text:
        # Simplemente reemplazamos el texto completo del párrafo
        para.text = para.text.replace(texto_antiguo, texto_nuevo)
        return True
    return False

# ============================================================================
# TAREA 1: Unificar cifras de productos
# ============================================================================
print("\n" + "="*80)
print("TAREA 1: Unificando cifras de productos...")
print("="*80)

for i, para in enumerate(doc.paragraphs):
    text = para.text
    
    # Reemplazar "400K+" o "400.000+" por la cifra real con contexto
    if "400K+ productos" in text:
        nuevo_texto = text.replace("400K+ productos", "~3.000 productos en Supabase (migración en curso desde los ~75.000 de Firestore)")
        para.text = nuevo_texto
        print(f"[{i}] 400K+ → ~3.000 productos (migración en curso)")
        cambios_realizados += 1
    
    elif "400.000+" in text and "producto" in text.lower():
        nuevo_texto = text.replace("400.000+", "~3.000 referencias (de ~75.000 en Firestore)")
        para.text = nuevo_texto
        print(f"[{i}] 400.000+ → ~3.000 referencias")
        cambios_realizados += 1
    
    # Asegurar que las menciones a 75K indiquen que es de Firestore
    if "75K" in text or "75.000" in text:
        if "Firestore" not in text and "productos" in text.lower():
            # Ya debería tener contexto pero verificamos
            pass

# ============================================================================
# TAREA 2: Corregir inconsistencias Firebase → Supabase
# ============================================================================
print("\n" + "="*80)
print("TAREA 2: Corrigiendo inconsistencias Firebase → Supabase...")
print("="*80)

for i, para in enumerate(doc.paragraphs):
    text = para.text
    
    # 4.1 Flujo 3: onAuthStateChanged → onAuthStateChange
    if "onAuthStateChanged" in text:
        para.text = text.replace("onAuthStateChanged", "onAuthStateChange (Supabase)")
        print(f"[{i}] onAuthStateChanged → onAuthStateChange")
        cambios_realizados += 1
    
    # 4.1 Flujo 3: Firebase Google OAuth → Supabase Google OAuth
    if "LoginPage → Firebase Google OAuth" in text:
        para.text = text.replace("LoginPage → Firebase Google OAuth", "LoginPage → Supabase Google OAuth")
        print(f"[{i}] LoginPage → Supabase Google OAuth")
        cambios_realizados += 1
    
    # 4.2 Stack: Firebase Auth → Supabase Auth
    if "Auth: Firebase Auth" in text and "Base de datos:" not in text:
        para.text = text.replace("Auth: Firebase Auth", "Auth: Supabase Auth")
        print(f"[{i}] Auth: Firebase Auth → Supabase Auth")
        cambios_realizados += 1
    
    # 4.2 Stack: Base de datos: Firestore (→ Supabase en) → Base de datos: Supabase (PostgreSQL)
    if "Base de datos: Firestore (→ Supabase en)" in text:
        para.text = text.replace("Base de datos: Firestore (→ Supabase en)", "Base de datos: Supabase (PostgreSQL)")
        print(f"[{i}] Base de datos: Firestore → Supabase (PostgreSQL)")
        cambios_realizados += 1
    
    # 4.4 Modelo de datos: "Actual: Firestore" → "Anterior: Firestore"
    if "Actual: Firestore con sincronización" in text:
        para.text = text.replace("Actual: Firestore con sincronización", "Anterior: Firestore (fase inicial)")
        print(f"[{i}] Actual: Firestore → Anterior: Firestore")
        cambios_realizados += 1
    
    # 4.4: "Este documento describe el modelo actual en Firestore" → "Este documento describe el modelo anterior"
    if "Este documento describe el modelo actual en Firestore" in text:
        para.text = text.replace("Este documento describe el modelo actual en Firestore", 
                                  "Este documento describe el modelo anterior (Firestore fue usado en la fase inicial)")
        print(f"[{i}] Modelo actual en Firestore → modelo anterior")
        cambios_realizados += 1
    
    # 4.4: "Modelo de datos en Supabase (futuro)" → "Modelo actual — Supabase (PostgreSQL)"
    if "Modelo de datos en Supabase (futuro)" in text:
        para.text = text.replace("Modelo de datos en Supabase (futuro)", "Modelo actual — Supabase (PostgreSQL)")
        print(f"[{i}] Supabase (futuro) → Supabase (actual)")
        cambios_realizados += 1
    
    # Eliminar "Estado: Migración en desarrollo" si aparece al final de secciones
    if "Estado: Migración en desarrollo" in text:
        para.text = text.replace("Estado: Migración en desarrollo", "").strip()
        print(f"[{i}] Eliminado 'Estado: Migración en desarrollo'")
        cambios_realizados += 1
    
    # Capa de seguridad: Firebase Security Rules → Supabase RLS
    if "Firebase Security Rules" in text:
        para.text = text.replace("Firebase Security Rules", "Supabase RLS (Row Level Security)")
        print(f"[{i}] Firebase Security Rules → Supabase RLS")
        cambios_realizados += 1
    
    # Tabla RNF-07: Firebase Auth → Spark → Supabase Auth → Free
    if "Firebase Auth → Spark" in text:
        para.text = text.replace("Firebase Auth → Spark", "Supabase Auth → Free")
        print(f"[{i}] Firebase Auth → Spark → Supabase Auth → Free")
        cambios_realizados += 1
    
    if "Firestore → Spark" in text:
        para.text = text.replace("Firestore → Spark", "Supabase → Free")
        print(f"[{i}] Firestore → Spark → Supabase → Free")
        cambios_realizados += 1
    
    # Capa 1: Firebase Auth → Supabase Auth
    if "Firebase Auth (Google Sign-In)" in text and "Capa" in text:
        para.text = text.replace("Firebase Auth (Google Sign-In)", "Supabase Auth (Google Sign-In)")
        print(f"[{i}] Capa 1: Firebase Auth → Supabase Auth")
        cambios_realizados += 1

# ============================================================================
# TAREA 3: Reparar frase rota del capítulo 5
# ============================================================================
print("\n" + "="*80)
print("TAREA 3: Reparando frase rota...")
print("="*80)

for i, para in enumerate(doc.paragraphs):
    text = para.text
    
    if "Este proyecto 15+ herramientas de IA" in text:
        # Frase completa según el prompt
        nuevo_texto = "En este proyecto se utilizaron más de 15 herramientas de IA distintas, desde Claude Web en marzo de 2026 hasta Hermes Agent en mayo de 2026. La metodología fue evolucionando conforme aparecían nuevas opciones:"
        para.text = nuevo_texto
        print(f"[{i}] Frase reparada")
        cambios_realizados += 1

# ============================================================================
# TAREA 4: Eliminar vestigio "Capítulo 02 — Estado del Arte"
# ============================================================================
print("\n" + "="*80)
print("TAREA 4: Eliminando vestigio 'Capítulo 02'...")
print("="*80)

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    if text == "*Capítulo 02 — Estado del Arte*" or text == "Capítulo 02 — Estado del Arte":
        para.text = ""  # Eliminar completamente
        print(f"[{i}] Eliminado 'Capítulo 02 — Estado del Arte'")
        cambios_realizados += 1

# ============================================================================
# TAREA 5: Arreglar estructura del documento (capítulo 1)
# ============================================================================
print("\n" + "="*80)
print("TAREA 5: Arreglando estructura del documento...")
print("="*80)

# Reescribir el bloque de estructura del documento
estructura_correcta = """Este TFC se documenta en 10 capítulos más una sección de anexos:

1. Introducción — Presentación del proyecto, objetivos y contexto
2. Análisis de la empresa — Sonepar: estructura, productos y contexto de las prácticas
3. Análisis de requisitos — Problemas identificados, requisitos funcionales y no funcionales
4. Diseño técnico — Arquitectura, stack tecnológico, UI/UX y modelo de datos
5. Contexto de la IA — Estado del arte de la IA generativa en desarrollo web
6. Herramientas de IA utilizadas — Catálogo de 15 herramientas con descripción y uso en el proyecto
7. Proceso de desarrollo — Metodología de trabajo con IA y cronología de las fases
8. Manuales de uso — Guías de los módulos principales
9. Resultados — Evaluación cualitativa y cuantitativa
10. Conclusiones y líneas futuras — Valoración final e impacto educativo
Anexos — Fichas técnicas completas de las herramientas de IA y tablas de referencia"""

encontrado_estructura = False
for i, para in enumerate(doc.paragraphs):
    text = para.text
    
    if "Estructura del documento" in text:
        encontrado_estructura = True
        # Los siguientes párrafos hasta el próximo capítulo son la estructura
        # Los reemplazamos con la estructura correcta
        print(f"[{i}] Encontrada 'Estructura del documento', reemplazando...")
        
        # Insertar la estructura correcta después de este párrafo
        para.text = text + "\n\n" + estructura_correcta
        
        # Eliminar los párrafos siguientes que contienen la estructura vieja
        # (hasta encontrar el próximo capítulo "2. ANÁLISIS")
        j = i + 1
        while j < len(doc.paragraphs):
            siguiente_text = doc.paragraphs[j].text.strip()
            if siguiente_text and (siguiente_text.startswith("2.") or "2. ANÁLISIS" in siguiente_text or "2. ANÁLISIS DE LA EMPRESA" in siguiente_text):
                break
            doc.paragraphs[j].text = ""  # Vaciar párrafos de la estructura vieja
            j += 1
        
        cambios_realizados += 1
        break

if not encontrado_estructura:
    print("AVISO: No se encontró 'Estructura del documento'")

# ============================================================================
# TAREA 10: Limpiar anglicismos ("Alternativas considered")
# ============================================================================
print("\n" + "="*80)
print("TAREA 10: Limpiando anglicismos...")
print("="*80)

for i, para in enumerate(doc.paragraphs):
    text = para.text
    
    if "Alternativas considered" in text:
        para.text = text.replace("Alternativas considered", "Alternativas consideradas")
        print(f"[{i}] 'Alternativas considered' → 'Alternativas consideradas'")
        cambios_realizados += 1
    
    # Corregir "depois" → "después" si aparece
    if "depois" in text.lower():
        para.text = text.replace("depois", "después")
        print(f"[{i}] 'depois' → 'después'")
        cambios_realizados += 1

# ============================================================================
# Guardar documento V2
# ============================================================================
print("\n" + "="*80)
print(f"Guardando documento V2 con {cambios_realizados} cambios...")
print("="*80)

doc.save(OUTPUT_PATH)
print(f"\nDocumento guardado en: {OUTPUT_PATH}")
print(f"Total cambios realizados: {cambios_realizados}")
print("\nNOTA: Las tareas 6, 7, 8 y 9 requieren modificaciones más complejas")
print("que se harán en un segundo paso.")