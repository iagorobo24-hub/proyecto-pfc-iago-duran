#!/usr/bin/env python3
"""
Genera MEMORIA_PFC_V4.docx — Versión 4 mejorada de la memoria del PFC Sonepar.
Mejoras sobre V3: métricas reales de código, schema SQL completo, arquitectura
documentada con datos reales, resultados del CTO audit, plan Phase 3, etc.
"""
import json, re, os, sys

# ── Add skill template path ──
sys.path.insert(0, '/home/abu/.hermes/skills/productivity/crear-documentos-academicos/templates')
from generate_academic_docx import (
    setup_document, add_cover, add_rich_para, add_word_table,
    group_into_blocks, parse_md_table, clean, strip_bold_markers
)
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

FONT  = 'Calibri'
BLUE  = RGBColor(0, 75, 141)
GRAY  = RGBColor(102, 102, 102)
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(0, 0, 0)
LTGRAY_BG = 'F5F7FA'

PROJECT_TITLE    = 'Suite de Herramientas Web para Técnicos del Sector Eléctrico'
PROJECT_SUBTITLE = 'Proyectos Sonepar — PFC 2025-2026'
PROJECT_TYPE     = 'PROYECTO FIN DE CICLO'
PROJECT_CYCLE    = 'Ciclo Formativo de Automatización y Robótica Industrial'

COVER_FIELDS = [
    ('Autor:',    'Iago Durán Romera'),
    ('Centro:',   'CIFP Universidade Laboral de A Coruña'),
    ('Tutor:',    'Jose Uzal'),
    ('Empresa:',  'Sonepar Ibérica'),
    ('Curso:',    '2025-2026'),
    ('Fecha:',    'Mayo 2026'),
]

CHAPTERS = [
    ('1', 'Resumen Ejecutivo',       'Visión general del proyecto y objetivos principales'),
    ('2', 'Estado del Arte',          'La IA Generativa como herramienta de desarrollo'),
    ('3', 'Análisis de Requisitos',   'Estudio de la empresa y necesidades identificadas'),
    ('4', 'Diseño Técnico',           'Arquitectura, stack tecnológico y modelo de datos'),
    ('5', 'Proceso de Desarrollo',    'Metodología, fases y lecciones aprendidas'),
    ('6', 'Herramientas de IA',       'Catálogo de herramientas utilizadas'),
    ('7', 'Manuales de Uso',          'Guías detalladas de los módulos principales'),
    ('8', 'Resultados',               'Evaluación cualitativa y cuantitativa del proyecto'),
    ('9', 'Conclusiones y Líneas Futuras', 'Valoración final e impacto educativo'),
    ('10','Manual para Profesores',    'Guía docente para replicar la metodología'),
    ('A', 'Anexos',                   'Tablas resumen del proyecto'),
]

# ──────────────────────────────────────────────────────────
# CONTENT — chapter by chapter
# ──────────────────────────────────────────────────────────
elements = []

def H1(n, title):
    elements.append({'type': 'heading1', 'text': f'{n}. {title}'})

def H2(title):
    elements.append({'type': 'heading2', 'text': title})

def H3(title):
    elements.append({'type': 'heading3', 'text': title})

def P(text):
    elements.append({'type': 'paragraph', 'text': text})

def B(text):
    elements.append({'type': 'bullet', 'text': text})

def TABLE(rows, caption=''):
    """rows: list of lists [header, row1, row2, ...]"""
    for r in rows:
        elements.append({'type': 'paragraph', 'text': '| ' + ' | '.join(str(c) for c in r) + ' |'})
    return

# ──────────────────────────────────────────
# PORTADA
# ──────────────────────────────────────────
P('')
P('SUITE DE HERRAMIENTAS WEB PARA TÉCNICOS DEL SECTOR ELÉCTRICO')
P('PROYECTO FIN DE CICLO')
P('Ciclo Formativo de Automatización y Robótica Industrial')
P('Autor: Iago Durán Romera')
P('Centro: CIFP Universidade Laboral de A Coruña')
P('Tutor: Jose Uzal')
P('Empresa: Sonepar Ibérica')
P('Curso: 2025-2026')
P('Fecha: Mayo 2026')
P('')

# ──────────────────────────────────────────
# ÍNDICE
# ──────────────────────────────────────────
H1('', 'ÍNDICE GENERAL')
for num, title, desc in CHAPTERS:
    P(f'{num}. {title} — {desc}' if desc else f'{num}. {title}')
P('')

# ══════════════════════════════════════════
# CAPÍTULO 1 — RESUMEN EJECUTIVO
# ══════════════════════════════════════════
H1('1', 'RESUMEN EJECUTIVO')
H2('1.1 — El proyecto en una frase')
P('Desarrollo de una suite de herramientas web para técnicos del sector eléctrico, de libre acceso y código abierto en GitHub, impulsada por IA generativa y desplegada en Vercel.')
H2('1.2 — ¿Qué es este proyecto?')
P('Es una aplicación web SPA (Single Page Application) con 7 módulos funcionales que resuelven necesidades reales de trabajadores relacionados al sector eléctrico: acceso a catálogos de productos, cálculo de presupuestos, simulación de almacén, gestión de incidencias, seguimiento de formación, asistente técnico con IA (SONEX) y dashboard de KPIs logísticos.')
H2('1.3 — Módulos creados')
TABLE([
    ['Módulo', 'Función', 'Tecnología principal'],
    ['Fichas Técnicas', 'Catálogo de productos eléctricos con navegación jerárquica', 'React 19 + Supabase + React Router v7'],
    ['SONEX', 'Asistente técnico con IA generativa en tiempo real', 'OpenRouter + Anthropic Claude + Framer Motion'],
    ['Simulador Almacén', 'Simulación de procesos logísticos (recepción, almacenamiento, picking, expedición)', 'React hooks + CSS Modules'],
    ['Presupuestos', 'Generador de presupuestos con exportador PDF', 'jsPDF + html2canvas + Supabase'],
    ['Formación Interna', 'Gestión de planes de formación por empleado', 'React + localStorage (sin backend)'],
    ['Incidencias', 'Registro y diagnóstico de incidencias con IA', 'React + Supabase + anthropicService'],
    ['KPI Logístico', 'Dashboards con métricas visuales', 'Recharts'],
])
P('')
H2('1.4 — Impacto cuantificable')
TABLE([
    ['Métrica', 'Valor'],
    ['Herramientas funcionales', '7 módulos'],
    ['Referencias de productos', '400.000+ scrapeadas de fuentes reales'],
    ['Líneas de código (src/)', '15.721 en 123 archivos'],
    ['Test unitarios', '272 tests pasando (12 archivos)'],
    ['Categorías de productos', '8 familias con navegación jerárquica'],
    ['Tiempo de build', '~9,7 segundos'],
    ['Coste total de infraestructura', '0 € (todo en tier gratuito)'],
    ['Tiempo de desarrollo', '~6 meses (Marzo–Mayo 2026)'],
    ['Vocabulary de herramientas IA documentadas', '15+'],
])
P('')
H2('1.5 — Objetivos alcanzados')
B('Catálogo real con más de 400.000 productos de fuentes auténticas (Schneider Electric, ABB, Siemens, Hager, Legrand, etc.)')
B('Arquitectura moderna documentada: React 19, Vite 7, Supabase, PWA, graceful degradation')
B('Migración de Firebase a Supabase completada con Row Level Security (RLS) activo')
B('Suite de tests unitarios con 272 tests cubriendo servicios, hooks y utilidades')
B('Metodología de desarrollo con IA generativa documentada paso a paso')
B('CTO-level audit con 31 hallazgos categorizados y plan de mejora Phase 3 priorizado')
B('Documentación completa: 10 capítulos + manual para profesores + anexos')
P('')

# ══════════════════════════════════════════
# CAPÍTULO 2 — ESTADO DEL ARTE
# ══════════════════════════════════════════
H1('2', 'ESTADO DEL ARTE')
H2('2.1 — La IA Generativa en el Desarrollo Web')
P('El desarrollo web con asistencia de IA generativa ha pasado por varias etapas diferenciadas:')
B('2020-2022 — Los orígenes: Primeros modelos de lenguaje (GPT-2, GPT-3) permiten generación de texto. Los desarrolladores experimentan con APIs de completado.')
B('2023 — Explosión: ChatGPT (Nov 2022) y Copilot (2021) popularizan el concepto. Se publican miles de artículos sobre "cómo usar IA para programar".')
B('2024 — Agentes y CLI: Aparecen agentes de IA que pueden ejecutar acciones (OpenCode, Claude CLI, Gemini CLI, Qwen CLI). Los modelos NVIDIA Nemotron producen código de alta calidad rivalizando con modelos propietarios.')
B('2025-2026 — Estado actual: La IA generativa se integra en el flujo diario de desarrollo. Herramientas como Claude Web, Windsurf, Cursor y GitHub Copilot son recursos estándar en equipos profesionales y académicos.')
H2('2.2 — Ventajas y riesgos documentados')
TABLE([
    ['Ventaja', 'Riesgo'],
    ['Productividad: automatiza tareas repetitivas (boilerplate, tests)', 'Calidad variable: el código generado puede tener errores sutiles'],
    ['Accesibilidad: reduce la barrera de entrada a la programación', 'Dependencia: usar IA sin entender el código impide el aprendizaje'],
    ['Aprendizaje: ver código generado ayuda a entender patrones', 'Seguridad: pueden generar código con vulnerabilidades si no se revisa'],
    ['Iteración rápida: versiones frecuentes con bajo coste', 'Obsolescencia: herramientas cambian rápidamente'],
])
P('')

# ══════════════════════════════════════════
# CAPÍTULO 3 — ANÁLISIS DE REQUISITOS
# ══════════════════════════════════════════
H1('3', 'ANÁLISIS DE REQUISITOS')
H2('3.1 — Análisis de la empresa')
P('El proyecto se desarrolla en colaboración con Sonepar Ibérica, líder mundial en la distribución de material eléctrico. El contexto de la empresa determina los requisitos funcionales de la aplicación.')
H2('3.2 — Problemas identificados')
B('Acceso a información de productos: Los técnicos necesitan能找到 información detallada de productos eléctricos (referencias, especificaciones, compatibilidades) sin acceso a catálogos impresos.')
B('Compatibilidad entre productos: Determinar si dos productos son compatibles entre sí requiere experiencia que no siempre está disponible.')
B('Cálculo de presupuestos: Los comerciales necesitan herramientas ágiles para calcular presupuestos sin depender de sistemas ERP complejos.')
B('Seguimiento de formación: La empresa necesita tracking del estado de formación de los empleados en normativas y procedimientos.')
B('Registro de incidencias: Los técnicos de almacén necesitan registrar y diagnosticar incidencias de forma estructurada.')
H2('3.3 — Requisitos funcionales principales')
TABLE([
    ['Código', 'Descripción', 'Prioridad', 'Módulo'],
    ['RF-01.1', 'Navegación jerárquica por familia → marca → gama → tipo', 'Alta', 'Fichas Técnicas'],
    ['RF-01.2', 'Búsqueda por referencia de fabricante o nombre', 'Alta', 'Fichas Técnicas'],
    ['RF-01.3', 'Ver ficha completa del producto con PDF y datos técnicos', 'Alta', 'Fichas Técnicas'],
    ['RF-02.1', 'Chat conversacional con IA técnica en tiempo real', 'Alta', 'SONEX'],
    ['RF-02.2', 'Streaming de respuestas (tokens aparecen progresivamente)', 'Alta', 'SONEX'],
    ['RF-02.3', 'Extracción de referencias de la respuesta IA con botones "Ver ficha"', 'Alta', 'SONEX'],
    ['RF-03.1', 'Wizard de 5 pasos para crear presupuestos', 'Alta', 'Presupuestos'],
    ['RF-03.2', 'Exportación a PDF con formato profesional', 'Alta', 'Presupuestos'],
    ['RF-03.3', 'Guardado automático en Supabase por usuario', 'Media', 'Presupuestos'],
    ['RF-04.1', 'Dashboard KPIs con gráficos (OEE, MTTR, MTBF, fill rate)', 'Alta', 'KPI'],
    ['RF-05.1', 'Simulación guiada de 4 etapas (recepción, almacenamiento, picking, expedición)', 'Alta', 'Simulador'],
    ['RF-06.1', 'CRUD de incidencias con AI diagnóstico automático', 'Alta', 'Incidencias'],
    ['RF-07.1', 'Matriz de formación por empleado y módulo', 'Media', 'Formación'],
    ['RF-08.1', 'Sistema de autenticación con Google OAuth via Supabase', 'Alta', 'Auth'],
    ['RF-08.2', 'Persistencia de preferencias por usuario (tema, datos)', 'Media', 'Auth'],
])
P('')
H2('3.4 — Requisitos no funcionales')
TABLE([
    ['RNF', 'Descripción', 'Objetivo'],
    ['RNF-01', 'Responsive design', '640-1024px (tablet/mobile)'],
    ['RNF-02', 'PWA instalable', 'Service Worker con workbox, manifest completo'],
    ['RNF-03', 'Tiempo de build', '< 10 segundos en Vercel (caché inteligente)'],
    ['RNF-04', 'Graceful degradation', 'App funciona (sin BD) si faltan variables de entorno'],
    ['RNF-05', 'Tests覆盖率', 'Core services y hooks con 272 tests unitarios'],
    ['RNF-06', 'Coste de infraestructura', '0 € (tier gratuito de todas las herramientas)'],
])
P('')

# ══════════════════════════════════════════
# CAPÍTULO 4 — DISEÑO TÉCNICO
# ══════════════════════════════════════════
H1('4', 'DISEÑO TÉCNICO')
H2('4.1 — Arquitectura del sistema')
P('La aplicación sigue una arquitectura SPA (Single Page Application) con React 19 y enrutamiento cliente-side via React Router v7. El backend de datos es Supabase (PostgreSQL + Auth + Realtime) y las llamadas a IA se proxean a través de /api/ai para no exponer claves API en el cliente.')
H2('4.2 — Stack tecnológico completo')
TABLE([
    ['Capa', 'Tecnología', 'Versión', 'Propósito'],
    ['Runtime', 'Node.js', '^20', 'Ejecución del entorno de desarrollo'],
    ['UI Framework', 'React', '^19.2.0', 'Componentes y estado de la interfaz'],
    ['Bundler', 'Vite', '^7.3.1', 'Build y desarrollo rápido'],
    ['PWA', 'vite-plugin-pwa', '^1.3.0', 'Service Worker, manifest, offline'],
    ['Routing', 'React Router DOM', '^7.13.1', 'Navegación cliente-side'],
    ['Backend', 'Supabase', '@supabase/supabase-js ^2.105.4', 'PostgreSQL + Auth + Realtime'],
    ['Auth', 'Supabase Auth (Google OAuth)', '—', 'Login con cuenta Google'],
    ['IA', 'OpenRouter / Anthropic Claude', '—', 'API proxy /api/ai en Vercel'],
    ['Animación', 'Framer Motion', '^12.38.0', 'Transiciones y microinteracciones'],
    ['Gráficos', 'Recharts', '^3.8.0', 'Dashboards KPI'],
    ['PDF', 'jsPDF + html2canvas', '^4.2.1 / ^1.4.1', 'Generación de documentos'],
    ['Scrapping', 'Crawlee + Camoufox + Puppeteer', '—', 'Scraping de catálogos'],
    ['Tests unitarios', 'Vitest + Testing Library', '^4.1.7 / ^16.3.2', 'Test coverage'],
    ['Tests E2E', 'Playwright', '^1.60.0', 'Tests de integración'],
    ['Linting', 'ESLint (flat config)', '^9.39.1', 'Code quality'],
])
P('')
H2('4.3 — Modelo de datos (Supabase — PostgreSQL)')
P('Esquema de las 4 tablas implementadas con Row Level Security (RLS):')
H3('Tabla: products (catálogo de productos)')
P('| Columna | Tipo | Notas |')
P('|---|---|---|')
P('| id | BIGSERIAL | PRIMARY KEY |')
P('| ref_fabricante | TEXT | INDEX, búsqueda por referencia |')
P('| name | TEXT | Nombre comercial del producto |')
P('| marca | TEXT | INDEX, nombre de marca (Schneider, ABB...) |')
P('| familia | TEXT | INDEX, familia de producto (CABLES, POTENCIA...) |')
P('| subfamilia | TEXT | INDEX |')
P('| tipo | TEXT | INDEX |')
P('| Gama | TEXT | INDEX |')
P('| Subgama | TEXT | INDEX |')
P('| precio | NUMERIC | Precio en euros |')
P('| imagen | TEXT | URL de imagen del producto |')
P('| pdf_url | TEXT | URL del datasheet PDF |')
P('| brand_id | INTEGER | FK → brands.id |')
P('')
H3('Tabla: brands')
P('| Columna | Tipo | Notas |')
P('|---|---|---|')
P('| id | SERIAL | PRIMARY KEY |')
P('| name | TEXT | Nombre oficial de marca |')
P('| website_url | TEXT | URL de la marca |')
P('')
H3('Tabla: testimonios (públicos, sin auth)')
P('| Columna | Tipo | Notas |')
P('|---|---|---|')
P('| id | BIGSERIAL | PRIMARY KEY |')
P('| user_id | UUID | FK → auth.users(id) ON DELETE SET NULL |')
P('| nombre | TEXT | NOT NULL |')
P('| email | TEXT | |')
P('| texto | TEXT | NOT NULL |')
P('| rating | INTEGER | CHECK 1-5 |')
P('| created_at | TIMESTAMPTZ | DEFAULT NOW() |')
P('RLS: SELECT público, INSERT público, UPDATE/DELETE solo autor')
P('')
H3('Tabla: user_data (clave-valor genérico por usuario)')
P('| Columna | Tipo | Notas |')
P('|---|---|---|')
P('| id | BIGSERIAL | PRIMARY KEY |')
P('| user_id | UUID | FK → auth.users(id) ON DELETE CASCADE |')
P('| module | VARCHAR(50) | Módulo (tema, preferencias, etc.) |')
P('| key | VARCHAR(100) | Clave específica |')
P('| data | JSONB | Datos estructurados |')
P('| created_at | TIMESTAMPTZ | DEFAULT NOW() |')
P('| updated_at | TIMESTAMPTZ | DEFAULT NOW() |')
P('UNIQUE(user_id, module, key)')
P('RLS: solo el propio usuario puede CRUD')
P('')
H2('4.4 — Servicios backend (src/services/)')
H3('catalogService.ts (626 líneas, TypeScript)')
P('Servicio principal de navegación del catálogo de productos. Carga familias, marcas, gamas, tipos y productos desde Supabase con paginación (1000 registros/página). Implementa getCategorias(), getMarcasPorCategoria(), getGamasPorMarcaYCategoria(), getTiposPorGamaMarcaYFamilia(), buscarProductos(), getProductoPorRef(). Usa validación de productos via validate.ts y caché en memoria para marcas.')
H3('anthropicService.ts (195 líneas, TypeScript)')
P('Capa de integración con OpenRouter para llamadas a IA. Implementa callAnthropicAI() (no-streaming) y callAnthropicAIStream() (streaming con callbacks onChunk/onDone). Incluye rate limiting cliente-side: 20 llamadas/minuto con store en memoria. Sanitiza URLs antes de mostrarlas. Parsea respuestas JSON y markdown.')
H3('brandLogoService.js (178 líneas, JavaScript)')
P('Resolución de logos de marca. Tiene 38 mappings hardcoded a archivos PNG en /logos/. Si no existe logo, genera un avatar con gradiente usando los 2 primeros caracteres de la marca (o inicial de las 2 primeras palabras). Selección de gradiente determinista via hash del nombre para consistencia.')
H2('4.5 — Contextos de React (src/contexts/)')
TABLE([
    ['Nombre', 'Props/Estado', 'Propósito'],
    ['AuthContext.jsx (86L)', 'user, loading, loginWithGoogle(), logout()', 'Google OAuth via Supabase. Incluye window.__PW_MOCK_USER__ para E2E. Stub if no env vars.'],
    ['ThemeContext.jsx (128L)', 'theme, toggle(), systemPreference', 'Light/dark toggle con View Transitions API (animación circular desde click). Persiste en Supabase user_data + localStorage fallback.'],
    ['ToastContext.jsx (80L)', 'show(mensaje, tipo, duracion)', 'Sistema de notificaciones en stack fijo bottom-right. Tipos: info, success, error, warning. Auto-dismiss con setTimeout.'],
])
P('')
H2('4.6 — Hooks personalizados (src/hooks/)')
TABLE([
    ['Hook', 'Líneas', 'Propósito principal'],
    ['useNavegacionFichas.js', '715', 'Estado de navegación del catálogo: paso actual, categoría, marca, gama, tipo, subcategoría. 17 useState, 11 useEffect con deps parciales. Pendiente refactor a useReducer (CTO Phase 3).'],
    ['useProductTable.js', '352', 'Estado de tabla de productos: ordenamiento, filtrado, paginación.'],
    ['useSimuladorAlmacen.js', '324', 'Lógica de simulación de almacén: 4 etapas, escenarios de incidencias.'],
    ['useSimuladorMultijugador.js', '196', 'Sincronización multiplayer de la simulación (WebSocket/Supabase Realtime).'],
    ['useSonex.js', '204', 'Estado del chat SONEX: historial de mensajes, streaming de IA, refs a fichas.'],
    ['useUserData.js', '191', ' CRUD sobre tabla user_data de Supabase por módulo.'],
    ['useTestimonios.js', '174', 'Carga y envío de testimonios públicos. Fallback a localStorage.'],
    ['useMemoriaUsuario.js', '134', 'Sistema de memoria por usuario con esquema definido (fichas, presupuestos, sonex, simulador, incidencias, kpi, formación). Persiste en localStorage y Supabase.'],
    ['useFichasTecnicas.js', '132', 'Estado de la herramienta Fichas: detalle de producto, búsqueda, carga de PDF.'],
    ['usePresupuestos.js', '107', 'Estado del wizard de presupuestos (5 pasos).'],
    ['useAnalytics.js', '103', 'Tracking de eventos de uso de herramientas.'],
    ['useKeyboardShortcuts.js', '56', 'Atajos de teclado globales (Ctrl+K para buscar, etc.).'],
    ['useDocumentTitle.js', '47', 'Actualiza document.title con el nombre de la herramienta activa.'],
])
P('')
H2('4.7 — Diseño UI/UX')
H3('Sistema de color (variables CSS en variables.css)')
TABLE([
    ['Variable', 'Valor', 'Uso'],
    ['--color-primary', '#004B8D', 'Color corporativo principal (botones, links, headings)'],
    ['--color-primary-light', '#4A90D9', 'Acentos, hover states'],
    ['--color-success', '#28A745', 'Indicadores de éxito, badges positivos'],
    ['--color-warning', '#FFC107', 'Alertas, warnings'],
    ['--color-error', '#DC3545', 'Errores, badges negativos'],
    ['--color-bg', '#F5F7FA', 'Fondo principal (light mode)'],
    ['--color-dark-bg', '#1A1A2E', 'Fondo principal (dark mode)'],
    ['--color-surface', '#FFFFFF / #252542', 'Tarjetas y componentes (light/dark)'],
])
P('')
H3('Tipografía')
TABLE([
    ['Tamaño', 'Uso'],
    ['24px', 'Títulos de página (h1)'],
    ['20px', 'Subtítulos (h2)'],
    ['18px', 'Títulos de sección (h3)'],
    ['16px', 'Texto de cuerpo principal'],
    ['14px', 'Texto secundario, labels'],
    ['12px', 'Captions, metadata, timestamps'],
])
P('')
H2('4.8 — Estrategia de code-splitting y rendimiento')
TABLE([
    ['Chunk', 'Contenido', 'Tamaño (gzip)'],
    ['vendor-react', 'react + react-dom + react-router', '73.71 kB'],
    ['vendor-animations', 'framer-motion + animate.css', '43.59 kB'],
    ['vendor-charts', 'recharts', '103.87 kB'],
    ['vendor-icons', 'lucide-react', '4.89 kB'],
    ['vendor-utils', 'dompurify + marked', '22.13 kB'],
    ['vendor-pdf', 'jspdf + html2canvas', '174.44 kB (NO preloaded)'],
    ['FichasTecnicas', 'Chunk lazily loaded por ruta', '19.23 kB'],
    ['Sonex', 'Chunk lazily loaded por ruta', '6.24 kB'],
])
P('')
P('Build: ~9.7 segundos, 69 entradas PWA precacheadas, 3.455 KB totales en dist/. Todos los chunks de vendor son controlados por manualChunks en vite.config.js. El chunk vendor-pdf (174 KB gzip) NO se pre-carrega ya que PDF solo se necesita en el módulo de presupuestos.')
H2('4.9 — PWA (Progressive Web App)')
TABLE([
    ['Configuración', 'Valor'],
    ['Plugin', 'vite-plugin-pwa ^1.3.0'],
    ['Register type', 'autoUpdate'],
    ['Manifest name', 'Proyectos PFC Tools'],
    ['Display', 'standalone'],
    ['Theme color', '#0072CE'],
    ['Service Worker', 'workbox con globPatterns para js/css/html/ico/svg/png/jpg/woff2'],
    ['Runtime caching', 'Google Fonts con CacheFirst y expiración 1 año (31536000s)'],
    ['Precache entries', '69 archivos (3455 KB)'],
    ['Icons', 'SVG 192x192 y 512x512 en /icons/'],
    ['Offline', 'Funciona con Service Worker una vez precacheados los assets'],
])
P('')

# ══════════════════════════════════════════
# CAPÍTULO 5 — PROCESO DE DESARROLLO
# ══════════════════════════════════════════
H1('5', 'PROCESO DE DESARROLLO')
H2('5.1 — Metodología: Cómo se trabajó con IA Generativa')
H3('Fase 1 — Prototipo inicial (artefactos sueltos)')
P('Los primeros componentes fueron archivos JSX independientes (sonepar-almacen-simulador.jsx, sonepar-fichas-tecnicas.jsx, etc.) abertos directamente en el navegador con scripts de CDN. Esta aproximación permitió validar ideas muy rápido sin overhead de build.')
H3('Fase 2 — Migración a Vite + React')
P('Una vez validados los artefactos, se migraron a una SPA con Vite y React Router. Esta fase incluyó: configuración de build con PWA, migrar de Firebase a Supabase, y separar cada herramienta como componente lazy-loaded.')
H3('Fase 3 — Estabilización')
P('Se implementaron contextos (Auth, Theme, Toast), se añadieron tests unitarios con Vitest, se configuró Supabase RLS, y se arreglaron stale closures y deps incompletas de useEffect detectadas durante el CTO audit.')
H3('Fase 4 — CTO Audit y Phase 3 Plan')
P('Se realizó un audit completo de código (15.721 LOC, 31 hallazgos) categorizados en Critical/High/Medium/Low. Se generó un plan de mejora priorizado para continuar el desarrollo.')
H2('5.2 — Workflow de prompting documentado')
B('Entender antes de pedir: Leer el código existente antes de pedir a la IA que genere.')
B('Pedir específico, no vago: Definir exactamente inputs, outputs, tecnología y restricciones.')
B('Revisar siempre: Verificar el código generado antes de commit.')
B('Iterar rápido: Si no funciona, explicar el error y pedir corrección.')
B('Documentar después: Actualizar EVOLUCIÓN.md con los cambios realizados.')
H2('5.3 — Lecciones aprendidas — Lo que funcionó')
B('Empezar simple: Los primeros componentes fueron artefactos JSX independientes, no una SPA completa.')
B('Iterar rápido: Cada día había una versión nueva funcional, aunque imperfecta.')
B('Documentar mientras trabajas: EVOLUCION.md se actualizaba en cada sesión.')
B('Usar herramientas gratuitas: Nunca necesitas pagar por herramientas de desarrollo.')
H2('5.4 — Lo que se podría mejorar (Phase 3)')
B('Tests desde el principio: Añadir tests E2E con Playwright antes de refactorizar.')
B('Validación con usuarios reales: Probar la aplicación con técnicos de Sonepar durante el desarrollo.')
B('Integración continua: GitHub Actions para build y tests automáticos.')
B('Arquitectura más robusta: useReducer para useNavegacionFichas (715 líneas con 17 useState).')

# ══════════════════════════════════════════
# CAPÍTULO 6 — HERRAMIENTAS DE IA UTILIZADAS
# ══════════════════════════════════════════
H1('6', 'HERRAMIENTAS DE IA UTILIZADAS')
H2('6.1 — Claude Web (Anthropic)')
P('Uso principal: Generación de componentes JSX desde descripciones en lenguaje natural. Workflow: describir la herramienta → copiar código → npm run dev → iterar.')
P('Calificación: ★★★★★')
H2('6.2 — GitHub Copilot')
P('Uso principal: Autocompletado en VSCode. Aceptar sugerencias con Tab, rechazar y escribir manualmente si no coinciden.')
P('Calificación: ★★★★★')
H2('6.3 — Vercel')
P('Uso: Hosting del proyecto. Conecta con GitHub, detecta React+Vite automáticamente, soporta entorno serverless functions (/api/ai proxy).')
P('Coste: 0 € (tier gratuito: 100GB transferencia/mes, 100h serverless/mes)')
P('Calificación: ★★★★★')
H2('6.4 — Windsurf IDE (Codeium)')
P('Uso: IDE con IA integrada. Autocompletado + Cascade AI (chat context-aware).')
P('Calificación: ★★★★☆')
H2('6.5 — Qwen CLI')
P('Uso: Ejecución de tareas en terminal via lenguaje natural. Ejecuta acciones en el filesystem.')
P('Calificación: ★★★★☆')
H2('6.6 — Gemini CLI')
P('Uso: Similar a Qwen. Ejecución de tareas desde terminal.')
P('Calificación: ★★★★★')
H2('6.7 — OpenCode CLI')
P('Uso: Agente de coding en terminal. Acceso a filesystem y ejecución de comandos.')
P('Calificación: ★★★★☆')
H2('6.8 — OpenRouter')
P('Uso: Aggregation layer para múltiples proveedores de IA (Anthropic, OpenAI, Google, etc.). Se usa como proxy en el backend de Vercel (/api/ai) para no exponer API keys en el cliente.')
P('Coste: 0 € con modelos gratuitos (anthropic/claude-3.5-haiku, google/gemini-2.0-flash, etc.)')
P('Calificación: ★★★★☆')
H2('6.9 — Supabase')
P('Uso: Backend completo (PostgreSQL + Auth + Realtime). Originally Firebase was used; migrated to Supabase for better SQL query capabilities and RLS.')
P('Coste: 0 € (tier gratuito: 500MB DB, 1GB storage, 50K usuarios)')
P('Calificación: ★★★★☆')

# ══════════════════════════════════════════
# CAPÍTULO 7 — MANUALES DE USO
# ══════════════════════════════════════════
H1('7', 'MANUALES DE USO')
H2('7.1 — Manual de Usuario: Simulador de Almacén')
P('El Simulador de Almacén permite practicar el flujo completo de gestión logística en 4 etapas:')
B('Recepción: Introduce los datos del pedido recibido y confirma la recepción.')
B('Almacenamiento: Asigna ubicación a cada producto y verifica la capacidad del almacén.')
B('Preparación de pedido: Introduce el pedido del cliente, el sistema verifica disponibilidad y confirma la preparación.')
B('Expedición: Confirma los productos preparados, genera el albarán y registra la salida.')
H2('7.2 — Manual de Usuario: SONEX — Asistente Técnico')
P('SONEX es un asistente de IA para consultas técnicas sobre mantenimiento industrial y productos eléctricos.')
B('Para hacer una pregunta: Escribe en el campo de texto y pulsa Enter. SONEX responde con streaming progresivo.')
B('Para ver una ficha de producto: Si la respuesta contiene referencias a productos, aparecen botones "Ver ficha" que abren la ficha técnica directamente.')
B('Para añadir a presupuesto: Desde la ficha del producto, haz clic en "Añadir a presupuesto".')
B('Para exportar la conversación: Haz clic en el botón "Exportar" para generar un PDF con toda la conversación.')
H2('7.3 — Manual de Usuario: Fichas Técnicas')
P('El navegador de fichas técnicas permite explorar el catálogo de productos de forma jerárquica:')
B('Navegación: Categoría → Marca → Gama → Tipo → Productos filtrados.')
B('Búsqueda directa: Escribe una referencia o nombre en la barra de búsqueda.')
B('Ver ficha: Haz clic en cualquier producto para ver su ficha completa con imagen, especificaciones y PDF.')
B('Añadir a presupuesto: Desde la ficha, pulsa "Añadir a presupuesto" para transfererir el producto al wizard.')
H2('7.4 — Manual de Usuario: Presupuestos')
P('El generador de presupuestos tiene 5 pasos:')
B('Paso 1: Seleccionar cliente y fecha.')
B('Paso 2: Elegir productos del catálogo o añadirlos manualmente.')
B('Paso 3: Revisar y ajustar cantidades y descuentos.')
B('Paso 4: Añadir observaciones y condiciones.')
B('Paso 5: Generar PDF y guardar en Supabase.')

# ══════════════════════════════════════════
# CAPÍTULO 8 — RESULTADOS
# ══════════════════════════════════════════
H1('8', 'RESULTADOS')
H2('8.1 — Resultados cualitativos')
TABLE([
    ['Módulo', 'Valor entregado'],
    ['Fichas Técnicas', 'Catálogo con 400K+ productos, navegación jerárquica de 8 familias, logo de marca automático'],
    ['SONEX', 'Asistente con streaming en tiempo real, extracción de referencias, diagnóstico IA'],
    ['Simulador Almacén', '4 etapas implementadas con escenarios de incidencias'],
    ['Presupuestos', 'Wizard de 5 pasos, exportador PDF con formato profesional'],
    ['Formación', 'Matriz por empleado y módulo, con KPIs de progreso'],
    ['Incidencias', 'CRUD completo con AI diagnóstico'],
    ['KPI Logístico', '6 KPIs con gráficos de semáforo (OEE, MTTR, MTBF, Fill Rate)'],
])
P('')
H2('8.2 — Resultados cuantitativos')
TABLE([
    ['Métrica', 'Valor', 'Notas'],
    ['Líneas de código', '15.721', 'src/ solo (123 archivos)'],
    ['Tests unitarios', '272 pasando', '12 archivos de test'],
    ['Build time', '9,7 segundos', 'En Vercel con caché'],
    ['Dist size', '3.455 KB', '69 entradas PWA'],
    ['Vendor chunks', '6', 'React, animations, charts, icons, utils, pdf'],
    ['Coste total', '0 €', 'Todas las herramientas en tier gratuito'],
    ['Categorías de producto', '8', 'Con navegación jerárquica completa'],
    ['Hooks personalizados', '13', '2.525 líneas totales'],
    ['Componentes UI', '21', 'Reutilizables en toda la app'],
    ['Servicios backend', '3', 'catalogService, anthropicService, brandLogoService'],
])
P('')
H2('8.3 — Valoración técnica')
TABLE([
    ['Aspecto', 'Valoración', 'Comentario'],
    ['Arquitectura', '7/10', 'Necesita useReducer para useNavegacionFichas (715L con 17 useState)'],
    ['Testing', '8/10', '272 tests cubriendo servicios y hooks core'],
    ['Documentación', '9/10', '10 capítulos, manual para profesores, CT0 audit'],
    ['Seguridad', '6/10', 'Auth bypass para tests (window.__PW_MOCK_USER__), 45 console.log en prod'],
    ['Rendimiento', '6/10', 'Vendor chunks OK pero hay 3 chunks >500KB sin code-split'],
    ['Code quality', '6.5/10', '31 hallazgos en CTO audit (4 críticos, 7 altos)'],
])
P('')
H2('8.4 — CTO Audit — Resumen de hallazgos (31 totales)')
TABLE([
    ['Severidad', 'Cantidad', 'Ejemplos'],
    ['CRITICAL', '4', 'useNavegacionFichas: 715L, 17 useState, 11 useEffect con deps incompletas. Stale closure en volver(). Auth bypass (window.__PW_MOCK_USER__).'],
    ['HIGH', '7', '8 dead exports en CircleLayout.jsx. 75+ inline styles. 10+ funciones inline sin useCallback en SimuladorAlmacen. Falta React.memo en componentes que reciben objetos.'],
    ['MEDIUM', '11', '45 console.log/warn/error en producción. useEffect deps incompletas en useSonex, useTestimonios. 2 vendor chunks >500KB (vendor-pdf 588KB, vendor-charts 347KB).'],
    ['LOW', '9', 'BEM inconsistente en CSS. Magic values (800ms, 250ms, 3s). document.title con window.location en SSR. Sin ErrorBoundary en rutas.'],
])
P('')
P('Informe completo en: .hermes/plans/fase-3-audit-report.md')

# ══════════════════════════════════════════
# CAPÍTULO 9 — CONCLUSIONES
# ══════════════════════════════════════════
H1('9', 'CONCLUSIONES Y LÍNEAS FUTURAS')
H2('9.1 — ¿Se cumplieron los objetivos?')
TABLE([
    ['Objetivo', 'Estado'],
    ['Catálogo con 400K+ productos y navegación jerárquica', 'CUMPLIDO'],
    ['Asistente técnico con IA (SONEX)', 'CUMPLIDO'],
    ['Simulador de almacén con 4 etapas', 'CUMPLIDO'],
    ['Dashboard KPI con gráficos', 'CUMPLIDO'],
    ['Wizard de presupuestos con PDF', 'CUMPLIDO'],
    ['Formación interna con matriz por empleado', 'CUMPLIDO'],
    ['Gestión de incidencias con IA', 'CUMPLIDO'],
    ['Suite de tests (272 tests)', 'CUMPLIDO'],
    ['Documentación completa', 'CUMPLIDO'],
    ['Integración con sistemas SAP/ERP de Sonepar', 'NO CUMPLIDO'],
    ['APP móvil nativa', 'NO CUMPLIDO'],
    ['Validación con usuarios reales', 'NO CUMPLIDO'],
])
P('')
H2('9.2 — Líneas futuras (Phase 3 Plan)')
H3('CRITICAL — Implementar inmediatamente')
B('C1: Refactorizar useNavegacionFichas.js (715L, 17 useState) → useReducer con NavigationState reducer para eliminar stale closures y deps incompletas.')
B('C2: Crear FichasTecnicasContext para eliminar prop drilling de 35 props en FichasTecnicasContent.')
B('C3: Extraer StepReferenciasSimple subcomponente y eliminar IIFE anidada en StepReferencias.jsx (457L).')
B('C4: Eliminar auth bypass window.__PW_MOCK_USER__ de AuthContext y substituir con mock de Supabase Auth en tests E2E.')
H3('HIGH — Implementar en las próximas semanas')
B('H1: Eliminar 8 dead exports de CircleLayout.jsx, dead SkeletonCard de FichasTecnicasSkeleton.jsx, y dead usePresupuestosContext.')
B('H2: Migración parcial a TypeScript de los archivos más críticos (useNavegacionFichas, useProductTable, catalog types).')
B('H3: Envolver las 10+ funciones inline de SimuladorAlmacen.jsx con useCallback para evitar re-renders innecesarios.')
B('H4: Dynamic import() para vendor-pdf.js (588KB) y vendor-charts.js (347KB) que solo se cargan cuando se necesitan.')
B('H5: Limpiar useSimuladorAlmacen.js (324L, archivo no trackeado, 0 imports) — es un experimento que debe ser integrado o eliminado.')
B('H6: Añadir React.memo a los child components que reciben object/array props para evitar re-renders innecesarios.')
H3('MEDIUM — Siguiente prioridad')
B('M1: Eliminar los 45 console.log/warn/error de producción (wrapper dev-only).')
B('M2: Reemplazar los 75+ inline styles con CSS module classes para mantener coherencia visual.')
B('M3: Extraer StepReferenciasSimple a su propio archivo con nombre correcto.')
B('M4: Consolidar magic values (800ms, 250ms, 3s, 5s, 30s) en constantes con nombre descriptivo.')
H3('LOW — Si queda tiempo')
B('L1: Reemplazar window.location.reload() en App.jsx con React Router navigate().')
B('L2: Añadir ErrorBoundary a nivel de rutas.')
B('L3: Integración con Sonepar (API real si la empresa la proporciona).')
B('L4: APP móvil con React Native.')

# ══════════════════════════════════════════
# CAPÍTULO 10 — MANUAL PARA PROFESORES
# ══════════════════════════════════════════
H1('10', 'MANUAL PARA PROFESORES')
H2('10.1 — ¿Qué demuestra este proyecto?')
P('Este proyecto demuestra competencias del ciclo formativo de Automatización y Robótica Industrial en: desarrollo web full-stack, bases de datos relacionales, integración de APIs de IA, PWA, testing, deployment en producción, y documentación profesional.')
H2('10.2 — Las 5 herramientas esenciales')
TABLE([
    ['Herramienta', 'Uso en el proyecto', 'Alternativa'],
    ['Claude Web / Copilot', 'Generación de componentes React', 'Windsurf, Cursor'],
    ['Vercel', 'Hosting y CI/CD automático', 'Netlify, GitHub Pages'],
    ['Supabase', 'Backend (PostgreSQL + Auth)', 'Firebase, PlanetScale'],
    ['Vitest + Playwright', 'Tests unitarios y E2E', 'Jest + Cypress'],
    ['React Router v7', 'Navegación SPA', 'TanStack Router'],
])
P('')
H2('10.3 — Plan de actividades propuesto (4 sesiones de 3h)')
H3('Sesión 1: Introducción a la IA como herramienta de desarrollo')
B('Qué es un LLM y cómo funciona [15 min]')
B('Diferencia entre chat web y herramientas de coding [15 min]')
B('Actividad: Primer contacto con Claude Web — generar un componente React simple [45 min]')
H3('Sesión 2: Control de versiones y entorno de desarrollo')
B('Instalación de Windsurf [20 min]')
B('Configuración básica + conexión a GitHub [20 min]')
B('Conceptos de control de versiones (commit, push, branch, merge) [30 min]')
B('Actividad: Hacer un commit en GitHub desde Windsurf [45 min]')
H3('Sesión 3: Deployment y hosting')
B('Concepto de hosting y CDN [15 min]')
B('Registro en Vercel + conexión con GitHub [20 min]')
B('Qué es una Serverless Function [15 min]')
B('Actividad: Desplegar la calculadora de la Sesión 1 en Vercel [60 min]')
H3('Sesión 4: Proyecto integrador')
B('Diseñar 3 pantallas en papel o Figma [30 min]')
B('Pedir a Claude que genere el mockup HTML/CSS [45 min]')
B('Desplegar en Vercel y probar en el móvil [45 min]')
H2('10.4 — Rúbrica de evaluación')
TABLE([
    ['Criterio', '0 pts', '5 pts', '10 pts'],
    ['Uso de IA', 'Copia sin entender', 'Usa IA parcialmente', 'Usa IA como amplificador'],
    ['Calidad del código', 'No funciona o no compila', 'Funciona con warnings', 'Limpio, testeado, documentado'],
    ['Deployment', 'No despliega', 'Despliegue manual', 'CI/CD automático funcionando'],
    ['Documentación', 'Sin docs', 'README básico', 'Docs completas + decisiones de diseño'],
])
P('')
H2('10.5 — Recomendaciones para futuros proyectos')
B('Tiempo: Solo 3 meses de desarrollo real — scope realista desde el día 1.')
B('Validación: Probar la aplicación con usuarios reales antes de terminar.')
B('Tests: Empezar con tests desde la primera semana, no al final.')
B('Arquitectura: Diseñar antes de codificar — un diagrama de arquitectura ahorra mucho tiempo.')
B('IA como herramienta: Enseñar a usar IA, no a depender de ella. El objetivo es que el alumno entienda lo que genera la IA.')

# ══════════════════════════════════════════
# ANEXOS
# ══════════════════════════════════════════
H1('A', 'ANEXOS')
H2('A.1 — Estructura de directorios')
TABLE([
    ['Ruta', 'Descripción'],
    ['app/src/tools/', '7 herramientas principales (DashboardGlobal, FichasTecnicas, etc.)'],
    ['app/src/components/', '64 componentes organizados por dominio'],
    ['app/src/hooks/', '13 hooks personalizados (2.525 líneas totales)'],
    ['app/src/contexts/', '3 contextos (Auth, Theme, Toast)'],
    ['app/src/services/', '3 servicios (catalog, anthropic, brandLogo)'],
    ['app/src/supabase/', 'Cliente Supabase con stub de graceful degradation'],
    ['app/src/__tests__/', '12 archivos de test (272 tests)'],
    ['app/scripts/', 'Scrapers, migrations SQL, lib'],
    ['.hermes/plans/', 'CTO audit y plan Phase 3'],
])
P('')
H2('A.2 — Dependencias principales del proyecto')
TABLE([
    ['Dependencia', 'Versión', 'Propósito'],
    ['react', '^19.2.0', 'UI framework'],
    ['vite', '^7.3.1', 'Bundler'],
    ['react-router-dom', '^7.13.1', 'Routing'],
    ['@supabase/supabase-js', '^2.105.4', 'Backend'],
    ['framer-motion', '^12.38.0', 'Animación'],
    ['recharts', '^3.8.0', 'Gráficos'],
    ['jspdf', '^4.2.1', 'Generación PDF'],
    ['dompurify', '^3.4.5', 'Sanitización HTML'],
    ['vitest', '^4.1.7', 'Tests unitarios'],
    ['playwright', '^1.60.0', 'Tests E2E'],
])
P('')
H2('A.3 — API Routes implementadas')
TABLE([
    ['Ruta', 'Método', 'Propósito'],
    ['/api/ai', 'POST', 'Proxy serverless a OpenRouter — recibe body, devuelve stream o JSON'],
    ['/api/ai', 'OPTIONS', 'CORS preflight'],
])
P('')
H2('A.4 — Métricas de calidad de código (CTO Audit)')
TABLE([
    ['Métrica', 'Valor', 'Clasificación'],
    ['Total LOC (src/)', '15.721', '—'],
    ['Archivos fuente', '123', '—'],
    ['Componentes', '64', '—'],
    ['Hooks', '13', '—'],
    ['TypeScript (%)', '3.5%', 'Bajo — 4/123 archivos'],
    ['Hallazgos críticos', '4', 'Necesita atención inmediata'],
    ['Hallazgos altos', '7', 'Implementar en Phase 3'],
    ['Hallazgos medios', '11', 'Siguiente prioridad'],
    ['Hallazgos bajos', '9', 'Si queda tiempo'],
    ['Puntuación global', '6.5/10', 'Funcional con deuda técnica'],
])
P('')
H2('A.5 — Graceful Degradation — Diseño de stubs')
P('Cuando las variables de entorno VITE_SUPABASE_URL y VITE_SUPABASE_ANON_KEY no están definidas, el cliente Supabase retorna un stubProxy que permite que la aplicación continúe funcionando con datos vacíos (no crashea). Este diseño permite:')
B('Desarrollo local sin configuración de Supabase.')
B('Build en Vercel sin vars configuradas (muestra warnings en build log).')
B('La app carga y es navegable, pero BD muestra vacío.')
P('Este patrón es una decisión de diseño consciente: mejor app funcional sin BD que app rota.')
H2('A.6 — Licencia y reutilización')
P('Este proyecto y su documentación están bajo licencia MIT. Puedes:')
B('Usar las fichas de herramientas en clase.')
B('Adaptar el plan de actividades para otros ciclos formativos.')
B('Copiar la estructura de documentación para otros proyectos.')
B('Usar el código como referencia (con atribución).')
P('Este manual es un documento vivo. Si lo usas, por favor contribuye mejorando los materiales para los próximos alumnos.')

# ══════════════════════════════════════════
# GENERATE DOCUMENT
# ══════════════════════════════════════════
if __name__ == '__main__':
    output_path = '/home/abu/github_repos/proyecto-pfc-iago-duran/proyecto-fin-ciclo/desarrollo-entrega-final/MEMORIA_PFC_V4.docx'

    doc = setup_document()

    # ── Cover page ──
    for _ in range(4):
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    for text in [
        PROJECT_TITLE,
        PROJECT_SUBTITLE,
        PROJECT_TYPE,
        PROJECT_CYCLE,
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(22) if text == PROJECT_TITLE else Pt(14)
        r.font.color.rgb = BLUE if text == PROJECT_TITLE else GRAY

    for _ in range(2):
        doc.add_paragraph()

    for lbl, val in COVER_FIELDS:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f'{lbl} ')
        r1.bold = True
        r1.font.color.rgb = BLUE
        r2 = p.add_run(val)
        r2.font.size = Pt(11)

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)

    # ── Index ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('ÍNDICE GENERAL')
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = BLUE
    doc.add_paragraph()

    for num, title, desc in CHAPTERS:
        p = doc.add_paragraph()
        r1 = p.add_run(f'{num}. ' if num else '')
        r1.bold = True
        r1.font.color.rgb = BLUE
        r2 = p.add_run(title)
        r2.bold = True
        r2.font.size = Pt(12)
        if desc:
            r3 = p.add_run(f' — {desc}')
            r3.font.size = Pt(11)
            r3.italic = True
            r3.font.color.rgb = GRAY

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)

    # ── Content ──
    blocks = group_into_blocks(elements)
    table_counter = 0
    chapter_started = False

    for block_type, block_data in blocks:
        if block_type == 'md_table':
            parsed = parse_md_table(block_data)
            if len(parsed) >= 2:
                table_counter += 1
                add_word_table(doc, parsed, caption=f'Tabla {table_counter}.')
        else:
            for elem in block_data:
                etype = elem['type']
                text = clean(elem.get('text', ''))

                if etype.startswith('heading'):
                    level = int(etype[-1])
                    p = doc.add_paragraph(style=f'Heading {level}')
                    r = p.add_run(text)
                    r.font.color.rgb = BLUE

                elif etype == 'bullet':
                    if len(text) < 3:
                        continue
                    p = doc.add_paragraph(style='List Bullet')
                    parts = re.split(r'(\*\*.*?\*\*)', text)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            r = p.add_run(part[2:-2])
                            r.bold = True
                        else:
                            r = p.add_run(part)

                elif etype == 'paragraph':
                    add_rich_para(doc, text)

    doc.save(output_path)

    sz = os.path.getsize(output_path)
    words = sum(len(p.text.split()) for p in doc.paragraphs if p.text.strip())
    print(f'✅ Saved: {output_path}')
    print(f'   Paragraphs: {len(doc.paragraphs)}')
    print(f'   Tables: {len(doc.tables)}')
    print(f'   Words (approx): {words}')
    print(f'   Estimated pages: ~{words // 500}')
    print(f'   Size: {sz / 1024 / 1024:.1f} MB')