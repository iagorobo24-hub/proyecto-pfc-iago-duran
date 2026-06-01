#!/usr/bin/env python3
from docx import Document
import re
import sys

def process_text(text, replacements):
    original = text
    for pattern, replacement, context in replacements:
        text = re.sub(pattern, replacement, text)
    return text if text != original else None

def update_document(input_path, output_path):
    print(f"📂 Abriendo: {input_path}")
    doc = Document(input_path)
    
    replacements = [
        # Productos (todos los casos)
        (r'\b4\.000\+', '4.689', '4.000+ -> 4.689'),
        (r'\b400\.000\+', '4.689', '400.000+ -> 4.689'),
        (r'\b4\.000\s+productos', '4.689 productos', '4.000 productos -> 4.689'),
        
        # Tests (ambos casos)
        (r'(Tests\s+)200\+\s*\|?\s*.*✅\s*CUMPLIDO\s*\(\s*272\s*\)', r'\g<1>362 (272u + 90E2E)', 'Tests 200+ (272) -> 362'),
        (r'272\s+pasando', '362 (272u + 90E2E) pasando', '272 pasando -> 362'),
        (r'\b272\b', '362', '272 -> 362 (cualquier otro caso)'),
        
        # LOC
        (r'\b15\.721\b', '14.989', '15.721 LOC -> 14.989'),
    ]
    
    changes = 0
    
    # Párrafos
    for para in doc.paragraphs:
        if para.text:
            new_text = process_text(para.text, replacements)
            if new_text:
                para.text = new_text
                changes += 1
                print(f"  ✏️  Párrafo: {para.text[:80]}...")
    
    # Tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text:
                        new_text = process_text(para.text, replacements)
                        if new_text:
                            para.text = new_text
                            changes += 1
                            print(f"  ✏️  Tabla celda: {para.text[:80]}...")
    
    print(f"\n💾 Guardando: {output_path}")
    doc.save(output_path)
    
    print(f"\n✅ {changes} cambios aplicados en total")
    return changes

if __name__ == "__main__":
    inp = "/home/abu/github_repos/proyecto-pfc-iago-duran/proyecto-fin-ciclo/desarrollo-entrega-final/MEMORIA_PFC_V4.docx"
    out = "/home/abu/github_repos/proyecto-pfc-iago-duran/proyecto-fin-ciclo/desarrollo-entrega-final/MEMORIA_PFC_V4_FINAL_v2.docx"
    
    count = update_document(inp, out)
    print(f"\n🎉 Archivo final: {out}")
