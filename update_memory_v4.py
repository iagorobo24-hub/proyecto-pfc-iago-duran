#!/usr/bin/env python3
"""
Actualiza MEMORIA_PFC_V4.docx con los datos reales verificados:
- Productos: de "4.000+" o "400.000+" a "4.689"
- Tests: de "272" a "362" (con desglose en paréntesis)
- LOC: de "15.721" a "14.989"
"""

from docx import Document
import re
import sys

def replace_in_text(run, pattern, replacement):
    """Reemplaza texto en un run (fragmento de texto formateado)"""
    if run.text:
        new_text = re.sub(pattern, replacement, run.text)
        if new_text != run.text:
            run.text = new_text
            return True
    return False

def process_paragraph(para, replacements):
    """Procesa un párrafo buscando todos los patrones"""
    changes = []
    for run in para.runs:
        original_run_text = run.text
        modified_run_text = original_run_text
        
        for pattern, replacement, context in replacements:
            if modified_run_text:
                new_text = re.sub(pattern, replacement, modified_run_text)
                if new_text != modified_run_text:
                    changes.append(context)
                    modified_run_text = new_text
        
        if modified_run_text != original_run_text:
            run.text = modified_run_text
    
    return changes

def update_document(input_path, output_path):
    print(f"📂 Abriendo documento: {input_path}")
    doc = Document(input_path)
    
    # Listado de reemplazos
    # Nota: Los patrones deben ser precisos para no romper formato
    replacements = [
        # Productos
        (r'\b4\.000\+', '4.689', '4.000+ -> 4.689 (productos)'),
        (r'\b400\.000\+', '4.689', '400.000+ -> 4.689 (productos)'),
        (r'\b4\.000\s+productos\b', '4.689 productos', '4.000 productos -> 4.689'),
        
        # Tests - Reemplazo inteligente para contextos específicos
        # Caso 1: "Tests | 200+ | ✅ CUMPLIDO (272)"
        (r'(200\+\s*[|]?\s*.*✅\s*CUMPLIDO\s*\()272(\))', r'\g<1>362 (272u + 90E2E)\2', 'Tests 272 -> 362 desglose'),
        
        # Caso 2: "272 pasando" en tablas
        (r'272\s+pasando', '362 (272u + 90E2E) pasando', '272 pasando -> 362'),
        
        # Caso 3: "Tests 200+" suelto
        (r'Tests\s+200\+', 'Tests 362', 'Tests 200+ -> 362'),
        
        # LOC
        (r'\b15\.721\b', '14.989', '15.721 LOC -> 14.989'),
    ]
    
    all_changes = []
    
    # Procesar párrafos principales
    print("\n🔍 Escaneando párrafos...")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            changes = process_paragraph(para, replacements)
            if changes:
                all_changes.extend(changes)
                print(f"  ✏️  Cambios en párrafo {i+1}")
    
    # Procesar tablas (celdas contienen párrafos)
    print("\n🔍 Escaneando tablas...")
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                # Las celdas pueden tener múltiples párrafos
                for para in cell.paragraphs:
                    if para.text.strip():
                        changes = process_paragraph(para, replacements)
                        if changes:
                            all_changes.extend(changes)
                            print(f"  ✏️  Cambios en tabla {table_idx+1}, fila {row_idx+1}, celda {cell_idx+1}")
    
    # Guardar
    print(f"\n💾 Guardando en: {output_path}")
    doc.save(output_path)
    
    # Resumen
    print("\n" + "="*70)
    print("✅ DOCUMENTO ACTUALIZADO CON ÉXITO")
    print("="*70)
    
    # Agrupar cambios por tipo
    from collections import Counter
    change_counts = Counter(all_changes)
    
    print(f"\n📊 Total de reemplazos realizados: {len(all_changes)}")
    print("\nDetalles por tipo de cambio:")
    for change_type, count in change_counts.items():
        print(f"  • {change_type}: {count} veces")
    
    if not all_changes:
        print("\n⚠️  No se detectaron cambios. Posibles causas:")
        print("   1. El formato del documento es diferente al esperado")
        print("   2. Los números están en imágenes o gráficos")
        print("   3. Hay espacio/salineación extra entre dígitos")
        print("\n💡 Sugerencia: Haz los cambios manualmente en Word con Ctrl+H")
    
    return len(all_changes)

if __name__ == "__main__":
    input_file = "/home/abu/github_repos/proyecto-pfc-iago-duran/proyecto-fin-ciclo/desarrollo-entrega-final/MEMORIA_PFC_V4.docx"
    output_file = "/home/abu/github_repos/proyecto-pfc-iago-duran/proyecto-fin-ciclo/desarrollo-entrega-final/MEMORIA_PFC_V4_FINAL.docx"
    
    try:
        count = update_document(input_file, output_file)
        print(f"\n🎉 Proceso finalizado. Archivo guardado como:\n   {output_file}")
        if count > 0:
            print(f"   ✅ {count} cambios aplicados con éxito")
        else:
            print("   ⚠️  Ningún cambio detectado (verificar manualmente)")
        sys.exit(0)
    except Exception as e:
        print(f"\n❌ Error fatal: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)